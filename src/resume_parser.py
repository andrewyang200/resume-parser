# Complete code for src/resume_parser.py
# ----- START OF CODE BLOCK -----

import os
import json
import logging
import argparse
from typing import Dict, List, Optional, Tuple, Any
from pathlib import Path
import time
import datetime  # Added for timestamp

# --- Dependency Imports ---
from dotenv import load_dotenv
from openai import OpenAI, APIError, RateLimitError
import PyPDF2
import docx


# --- Configuration ---
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

# --- Constants ---
# Use the latest available high-quality model supporting JSON mode reliably.
# GPT-4o is generally recommended for complex parsing tasks.
DEFAULT_LLM_MODEL = "gpt-4o-mini"
MAX_RETRIES = 3
RETRY_DELAY_SECONDS = 5
MAX_TEXT_LENGTH = 20000 # Limit input text length to LLM to manage tokens/cost

# Define the target JSON structure we want the LLM to generate.
# This is critical for consistency.
JSON_SCHEMA_DEFINITION = """
{
  "schema_version": "1.0",
  "parsing_metadata": {
    "model_used": "string (e.g., gpt-4o)",
    "timestamp_utc": "string (ISO 8601 format)",
    "processing_time_seconds": "float"
  },
  "contact_information": {
    "name": "string | null",
    "email": "string | null",
    "phone": "string | null",
    "location": "string | null",
    "linkedin_url": "string | null",
    "portfolio_url": "string | null",
    "github_url": "string | null",
    "other_url": "string | null"
  },
  "summary": "string | null (Overall professional summary or objective)",
  "work_experience": [
    {
      "company": "string | null",
      "title": "string | null",
      "location": "string | null",
      "start_date": "string (YYYY-MM, YYYY, or original text) | null",
      "end_date": "string (YYYY-MM, YYYY, 'Present', or original text) | null",
      "duration": "string (Calculated or stated duration, e.g., '2 years 3 months') | null",
      "description": "string | null (Responsibilities, achievements, projects within the role)"
    }
  ],
  "education": [
    {
      "institution": "string | null",
      "degree": "string | null",
      "field_of_study": "string | null",
      "location": "string | null",
      "start_date": "string (YYYY-MM, YYYY, or original text) | null",
      "end_date": "string (YYYY-MM, YYYY, 'Present', or original text) | null",
      "graduation_date": "string (YYYY-MM, YYYY, or original text) | null",
      "description": "string | null (Minor, honors, thesis, relevant coursework)"
    }
  ],
  "skills": {
    "technical": ["string"],
    "soft": ["string"],
    "languages": ["string"],
    "certifications": ["string"],
    "other": ["string"]
  },
  "projects": [
    {
      "name": "string | null",
      "description": "string | null",
      "technologies_used": ["string"],
      "url": "string | null",
      "associated_experience": "string (e.g., company name or 'personal') | null"
    }
  ],
  "awards_and_recognition": [
    {
      "title": "string | null",
      "issuer": "string | null",
      "date": "string (YYYY-MM or YYYY) | null",
      "description": "string | null"
    }
  ],
  "publications": [
    {
        "title": "string | null",
        "authors": ["string"],
        "journal_or_conference": "string | null",
        "date": "string (YYYY-MM or YYYY) | null",
        "url_or_doi": "string | null"
    }
  ],
  "volunteer_experience": [
      {
          "organization": "string | null",
          "role": "string | null",
          "start_date": "string (YYYY-MM, YYYY, or original text) | null",
          "end_date": "string (YYYY-MM, YYYY, 'Present', or original text) | null",
          "description": "string | null"
      }
  ],
  "raw_text_preview": "string (First ~500 characters of extracted text)"
}
"""

# --- Helper Functions ---

def load_config() -> Optional[str]:
    """Loads OpenAI API key from .env file."""
    load_dotenv()
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        logging.error("OpenAI API key not found. Please set OPENAI_API_KEY in your .env file.")
        return None
    return api_key

def extract_text_from_pdf(file_path: Path) -> Optional[str]:
    """Extracts text content from a PDF file."""
    logging.info(f"Attempting to extract text from PDF: {file_path}")
    full_text = ""
    try:
        with open(file_path, 'rb') as pdf_file:
            reader = PyPDF2.PdfReader(pdf_file, strict=False) # strict=False might help with some malformed PDFs
            
            # Check if the PDF is encrypted
            if reader.is_encrypted:
                try:
                    # Attempt to decrypt with an empty password, common for unprotected files
                    # that still register as encrypted.
                    decrypt_result = reader.decrypt("")
                    if decrypt_result == 0: # 0 indicates failure to decrypt
                         logging.warning(f"PDF file {file_path.name} is encrypted and could not be decrypted with empty password.")
                         return None
                    elif decrypt_result == 1:
                         logging.info(f"PDF file {file_path.name} decrypted with empty password.")
                    elif decrypt_result == 2:
                         logging.info(f"PDF file {file_path.name} decrypted with provided password (though none was needed here).")

                except NotImplementedError:
                    logging.warning(f"Decryption for this PDF format ({file_path.name}) is not supported by PyPDF2.")
                    return None
                except Exception as decrypt_exc:
                    logging.error(f"Error during decryption attempt for {file_path.name}: {decrypt_exc}")
                    return None

            text_parts = []
            num_pages = len(reader.pages)
            logging.info(f"PDF has {num_pages} pages.")
            for i, page in enumerate(reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text_parts.append(page_text)
                    else:
                        logging.warning(f"No text extracted from page {i+1}/{num_pages} of {file_path.name}. It might be image-based or empty.")
                except Exception as page_exc:
                    logging.warning(f"Error extracting text from page {i+1}/{num_pages} of {file_path.name}: {page_exc}")

            # Join pages with double newline for better separation
            full_text = "\n\n".join(text_parts).strip()

            if not full_text:
                 logging.warning(f"No text extracted from PDF: {file_path.name}. It might be entirely image-based, corrupted, or encrypted without successful decryption.")
                 # Consider adding OCR as a fallback here if needed (outside MVP scope)
                 return None

            logging.info(f"Successfully extracted ~{len(full_text)} characters from PDF: {file_path.name}")
            return full_text

    except FileNotFoundError:
        logging.error(f"PDF file not found: {file_path}")
        return None
    except PyPDF2.errors.PdfReadError as e:
        logging.error(f"Error reading PDF file {file_path.name}: {e}. File might be corrupted or password-protected.")
        return None
    except Exception as e:
        logging.error(f"An unexpected error occurred during PDF extraction for {file_path.name}: {e}")
        # Log the full traceback for unexpected errors
        logging.exception("Detailed traceback for PDF extraction error:")
        return None


def extract_text_from_docx(file_path: Path) -> Optional[str]:
    """Extracts text content from a DOCX file."""
    logging.info(f"Attempting to extract text from DOCX: {file_path}")
    try:
        doc = docx.Document(file_path)
        text_parts = []
        for para in doc.paragraphs:
            text_parts.append(para.text)
        # Consider extracting text from tables as well for more completeness
        # for table in doc.tables:
        #     for row in table.rows:
        #         for cell in row.cells:
        #             text_parts.append(cell.text)

        full_text = "\n".join(text_parts).strip()
        if not full_text:
            logging.warning(f"No text extracted from DOCX: {file_path.name}. It might be empty.")
            return None
        logging.info(f"Successfully extracted ~{len(full_text)} characters from DOCX: {file_path.name}")
        return full_text
    except FileNotFoundError:
        logging.error(f"DOCX file not found: {file_path}")
        return None
    except Exception as e: # Catches errors from python-docx like invalid file format
        logging.error(f"Error reading DOCX file {file_path.name}: {e}")
        logging.exception("Detailed traceback for DOCX extraction error:")
        return None


def get_resume_text(file_path: Path) -> Optional[str]:
    """Determines file type and calls the appropriate text extraction function."""
    if not file_path.exists():
        logging.error(f"Input file does not exist: {file_path}")
        return None
    if not file_path.is_file():
        logging.error(f"Input path is not a file: {file_path}")
        return None

    suffix = file_path.suffix.lower()
    if suffix == '.pdf':
        return extract_text_from_pdf(file_path)
    elif suffix == '.docx':
        return extract_text_from_docx(file_path)
    else:
        logging.error(f"Unsupported file type: '{suffix}'. Please provide a PDF or DOCX file.")
        return None


# def parse_resume_with_llm(resume_text: str, client: OpenAI, model: str = DEFAULT_LLM_MODEL) -> Tuple[Optional[Dict[str, Any]], float]:
#     """Uses OpenAI's LLM to parse resume text into structured JSON. Returns parsed data and processing time."""
#     logging.info(f"Attempting to parse resume text using model: {model}")
#     start_time = time.time()

#     # Truncate text if it exceeds the defined limit
#     if len(resume_text) > MAX_TEXT_LENGTH:
#         logging.warning(f"Resume text length ({len(resume_text)}) exceeds limit ({MAX_TEXT_LENGTH}). Truncating.")
#         resume_text = resume_text[:MAX_TEXT_LENGTH]

#     system_prompt = f"""
# You are an expert Human Resources assistant specialized in accurately parsing résumés.
# Your task is to extract key information from the provided résumé text and structure it precisely according to the following JSON schema.
# Adhere strictly to the schema. Use `null` for fields where information is not found in the text.
# For lists (like work_experience, education, skills, projects), return an empty list `[]` if no relevant items are found.
# Do not add any information not present in the resume text. Do not invent or infer information beyond what is explicitly stated or clearly implied (e.g. calculate duration if start/end dates are present).
# Prioritize accuracy and completeness based *only* on the provided text.
# For dates, try to standardize to YYYY-MM or YYYY if possible, otherwise preserve the original text (e.g., "Spring 2020"). Use "Present" for current roles/education if indicated.
# Classify skills into the provided categories (technical, soft, languages, certifications, other) as accurately as possible based on common industry understanding. If unsure, place under 'other'.

# Target JSON Schema:
# ```json
# {JSON_SCHEMA_DEFINITION}

# Ensure your entire output is ONLY the valid JSON object, starting with {{ and ending with }}, without any introductory text, explanations, or markdown formatting like json ....
# """
    
#     user_prompt = f"""
# Please parse the following résumé text and return the extracted information as a valid JSON object adhering strictly to the schema provided in the system instructions.

# Résumé Text:
# --- START RESUME TEXT ---
# {resume_text}
# --- END RESUME TEXT ---

# Remember: Output ONLY the JSON object. Ensure all strings within the JSON are properly escaped.
# """
    
#     for attempt in range(MAX_RETRIES):
#       logging.info(f"LLM API call attempt {attempt + 1}/{MAX_RETRIES}")
#       try:
#           api_start_time = time.time()
#           response = client.chat.completions.create(
#               model=model,
#               messages=[
#                   {"role": "system", "content": system_prompt},
#                   {"role": "user", "content": user_prompt}
#               ],
#               temperature=0.0,  # Lower temperature for more deterministic parsing
#               response_format={"type": "json_object"}, # Use guaranteed JSON mode
#               # Consider adding max_tokens if needed, though JSON mode helps
#               # timeout=60 # Add a timeout for the API call
#           )
#           api_end_time = time.time()
#           logging.info(f"LLM API call successful (attempt {attempt+1}/{MAX_RETRIES}). API time: {api_end_time - api_start_time:.2f}s")

#           raw_response_content = response.choices[0].message.content
#           if not raw_response_content:
#               logging.error(f"LLM returned an empty response (attempt {attempt+1}).")
#               if attempt < MAX_RETRIES - 1:
#                   logging.info(f"Retrying in {RETRY_DELAY_SECONDS} seconds...")
#                   time.sleep(RETRY_DELAY_SECONDS)
#                   continue
#               else:
#                   return None, time.time() - start_time

#           # Attempt to parse the JSON content
#           try:
#               parsed_json = json.loads(raw_response_content)

#               # Basic validation: check if it's a dictionary (root level)
#               if not isinstance(parsed_json, dict):
#                   logging.error(f"LLM output was not a valid JSON object (root not a dict) (attempt {attempt+1}). Content: {raw_response_content[:500]}...")
#                   if attempt < MAX_RETRIES - 1:
#                       logging.info(f"Retrying in {RETRY_DELAY_SECONDS} seconds...")
#                       time.sleep(RETRY_DELAY_SECONDS)
#                       continue
#                   else:
#                       logging.error("Max retries reached for JSON format validation.")
#                       return None, time.time() - start_time

#               # Add parsing metadata dynamically
#               total_processing_time = time.time() - start_time
#               parsed_json.setdefault('parsing_metadata', {}) # Ensure key exists
#               parsed_json['parsing_metadata']['model_used'] = model
#               parsed_json['parsing_metadata']['timestamp_utc'] = datetime.datetime.now(datetime.timezone.utc).isoformat(timespec='seconds')
#               parsed_json['parsing_metadata']['processing_time_seconds'] = round(total_processing_time, 2)
#               # Add raw text preview safely
#               preview_text = resume_text[:500] + ('...' if len(resume_text) > 500 else '')
#               parsed_json['raw_text_preview'] = preview_text


#               logging.info("Successfully parsed LLM response into JSON.")
#               return parsed_json, total_processing_time # Return data and time

#           except json.JSONDecodeError as json_err:
#               logging.error(f"Failed to decode JSON from LLM response (attempt {attempt+1}). Error: {json_err}")
#               logging.error(f"Raw response content snippet: {raw_response_content[:500]}...") # Log beginning of problematic response
#               if attempt < MAX_RETRIES - 1:
#                   logging.info(f"Retrying in {RETRY_DELAY_SECONDS} seconds...")
#                   time.sleep(RETRY_DELAY_SECONDS)
#                   continue
#               else:
#                   logging.error("Max retries reached for JSON decoding.")
#                   return None, time.time() - start_time

#       except RateLimitError as e:
#           logging.warning(f"Rate limit exceeded (attempt {attempt+1}/{MAX_RETRIES}). Retrying after delay... Error: {e}")
#           if attempt < MAX_RETRIES - 1:
#               # Implement exponential backoff potentially
#               delay = RETRY_DELAY_SECONDS * (2 ** attempt)
#               logging.info(f"Retrying in {delay} seconds...")
#               time.sleep(delay)
#           else:
#               logging.error("Max retries reached due to rate limiting.")
#               return None, time.time() - start_time
#       except APIError as e:
#           # Handle specific API errors if needed (e.g., context length exceeded, auth errors)
#           logging.error(f"OpenAI API error occurred (attempt {attempt+1}/{MAX_RETRIES}): Status Code={e.status_code}, Message={e.message}")
#           if e.status_code == 401: # Unauthorized
#               logging.error("Authentication error. Check your OpenAI API key.")
#               return None, time.time() - start_time # Don't retry auth errors
#           if e.status_code == 429: # Should be caught by RateLimitError but as fallback
#               logging.warning("API reported rate limit error (APIError). Retrying...")
#               # Use similar backoff as RateLimitError
#               if attempt < MAX_RETRIES - 1:
#                   delay = RETRY_DELAY_SECONDS * (2 ** attempt)
#                   logging.info(f"Retrying in {delay} seconds...")
#                   time.sleep(delay)
#                   continue
#               else:
#                   logging.error("Max retries reached due to API rate limits.")
#                   return None, time.time() - start_time
#           # Handle other potentially retryable errors (e.g., 5xx server errors)
#           elif e.status_code >= 500:
#               if attempt < MAX_RETRIES - 1:
#                   logging.info(f"Retrying due to server error in {RETRY_DELAY_SECONDS} seconds...")
#                   time.sleep(RETRY_DELAY_SECONDS)
#                   continue
#               else:
#                   logging.error("Max retries reached due to API server errors.")
#                   return None, time.time() - start_time
#           else: # Non-retryable client-side errors
#               logging.error("Non-retryable API error encountered.")
#               return None, time.time() - start_time

#       except Exception as e:
#           logging.error(f"An unexpected error occurred during LLM parsing (attempt {attempt+1}/{MAX_RETRIES}): {type(e).__name__} - {e}")
#           logging.exception("Detailed traceback for LLM parsing error:")
#           # Generally, don't retry unknown errors unless certain they are transient
#           return None, time.time() - start_time

#     logging.error("Failed to parse resume after all retries.")
#     return None, time.time() - start_time

def parse_resume_with_llm_raw(resume_text: str, client: OpenAI, model: str = DEFAULT_LLM_MODEL) -> Tuple[Optional[str], float]:
    """
    Uses OpenAI's LLM to parse resume text into structured JSON string.
    Returns the raw JSON string from the LLM and processing time.
    Metadata (like model used, timestamp) should be added later in the pipeline.
    """
    logging.info(f"Attempting to parse resume text using model: {model}")
    start_time = time.time()

    if len(resume_text) > MAX_TEXT_LENGTH:
        logging.warning(f"Resume text length ({len(resume_text)}) exceeds limit ({MAX_TEXT_LENGTH}). Truncating.")
        resume_text = resume_text[:MAX_TEXT_LENGTH]

    system_prompt = f"""
      You are an expert Human Resources assistant specialized in accurately parsing résumés.
      Your task is to extract key information from the provided résumé text and structure it precisely according to the following JSON schema.
      Adhere strictly to the schema. Use `null` for fields where information is not found in the text.
      For lists (like work_experience, education, skills, projects), return an empty list `[]` if no relevant items are found.
      Do not add any information not present in the resume text. Do not invent or infer information beyond what is explicitly stated or clearly implied (e.g. calculate duration if start/end dates are present).
      Prioritize accuracy and completeness based *only* on the provided text.
      For dates, try to standardize to YYYY-MM or YYYY if possible, otherwise preserve the original text (e.g., "Spring 2020"). Use "Present" for current roles/education if indicated.
      Classify skills into the provided categories (technical, soft, languages, certifications, other) as accurately as possible based on common industry understanding. If unsure, place under 'other'.

      Target JSON Schema:
      ```json
      {JSON_SCHEMA_DEFINITION}

      Ensure your entire output is ONLY the valid JSON object, starting with {{ and ending with }}, without any introductory text, explanations, or markdown formatting like json ....
    """
    
    user_prompt = f"""
      Please parse the following résumé text and return the extracted information as a valid JSON object adhering strictly to the schema provided in the system instructions.

      Résumé Text:
      --- START RESUME TEXT ---
      {resume_text}
      --- END RESUME TEXT ---

      Remember: Output ONLY the JSON object. Ensure all strings within the JSON are properly escaped.
    """

    for attempt in range(MAX_RETRIES):
        logging.info(f"LLM API call attempt {attempt + 1}/{MAX_RETRIES}")
        try:
            api_start_time = time.time()
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.0,
                response_format={"type": "json_object"},
            )
            api_end_time = time.time()
            processing_time = time.time() - start_time
            logging.info(f"LLM API call successful (attempt {attempt+1}). API time: {api_end_time - api_start_time:.2f}s")

            raw_response_content = response.choices[0].message.content
            if not raw_response_content:
                logging.error(f"LLM returned an empty response (attempt {attempt+1}).")
                # Retry logic...
                if attempt < MAX_RETRIES - 1:
                    logging.info(f"Retrying in {RETRY_DELAY_SECONDS} seconds...")
                    time.sleep(RETRY_DELAY_SECONDS)
                    continue
                else:
                    return None, processing_time # Failed after retries

            # Basic check if it looks like JSON before returning
            if raw_response_content.strip().startswith('{') and raw_response_content.strip().endswith('}'):
                return raw_response_content, processing_time # Return raw string
            else:
                logging.error(f"LLM output doesn't look like JSON (attempt {attempt+1}). Content: {raw_response_content[:100]}...")
                # Retry logic...
                if attempt < MAX_RETRIES - 1:
                    logging.info(f"Retrying in {RETRY_DELAY_SECONDS} seconds...")
                    time.sleep(RETRY_DELAY_SECONDS)
                    continue
                else:
                    return None, processing_time # Failed after retries

        except RateLimitError as e:
            logging.warning(f"Rate limit exceeded (attempt {attempt+1}). Retrying...")
            # Retry logic with backoff...
            if attempt < MAX_RETRIES - 1:
                delay = RETRY_DELAY_SECONDS * (2 ** attempt)
                time.sleep(delay)
            else:
                logging.error("Max retries reached due to rate limiting.")
                return None, time.time() - start_time
        except APIError as e:
            logging.error(f"OpenAI API error (attempt {attempt+1}): {e}")
            # Retry logic for server errors (5xx)...
            if e.status_code >= 500 and attempt < MAX_RETRIES - 1:
                time.sleep(RETRY_DELAY_SECONDS)
                continue
            else: # Non-retryable or max retries reached
                return None, time.time() - start_time
        except Exception as e:
            logging.error(f"Unexpected error during LLM parsing (attempt {attempt+1}): {e}")
            logging.exception("Detailed traceback for LLM parsing error:")
            return None, time.time() - start_time # Don't retry unexpected

    logging.error("Failed to parse resume after all retries.")
    return None, time.time() - start_time


# def save_parsed_data(data: Dict[str, Any], output_path: Path) -> bool:
#     """Saves the parsed JSON data to a file."""
#     try:
#         output_path.parent.mkdir(parents=True, exist_ok=True)  # Ensure output directory exists
#         with open(output_path, 'w', encoding='utf-8') as f:
#             json.dump(data, f, indent=2, ensure_ascii=False)
#         logging.info(f"Successfully saved parsed data to: {output_path}")
#         return True
#     except IOError as e:
#         logging.error(f"Failed to write parsed data to {output_path}: {e}")
#         return False
#     except Exception as e:
#         logging.error(f"An unexpected error occurred during saving: {e}")
#         logging.exception("Detailed traceback for saving error:")
#         return False
    
# def main():
#     parser = argparse.ArgumentParser(
#         description="Parse a resume file (PDF or DOCX) into structured JSON using an LLM."
#     )
#     parser.add_argument(
#         "input_file",
#         help="Path to the input resume file (PDF or DOCX)."
#     )
#     parser.add_argument(
#         "-o",
#         "--output_file",
#         help="Path to save the output JSON file. Defaults to '<input_filename>.json' in 'data/parsed/'."
#     )
#     parser.add_argument(
#         "--model",
#         default=DEFAULT_LLM_MODEL,
#         help=f"OpenAI model to use for parsing (default: {DEFAULT_LLM_MODEL})."
#     )
#     parser.add_argument(
#         "--output_dir",
#         default="data/parsed",
#         help="Directory to save the output JSON files (default: data/parsed)."
#     )
#     args = parser.parse_args()

#     input_path = Path(args.input_file)

#     # Determine output path
#     if args.output_file:
#         output_path = Path(args.output_file)
#     else:
#         # Default output path construction
#         output_dir = Path(args.output_dir)
#         output_filename = input_path.stem + ".json"
#         output_path = output_dir / output_filename

#     logging.info(f"Starting resume parsing process for: {input_path}")
#     logging.info(f"Using LLM model: {args.model}")
#     logging.info(f"Output will be saved to: {output_path}")

#     overall_start_time = time.time()

#     # 1. Load API Key
#     api_key = load_config()
#     if not api_key:
#         return # Error already logged

#     # Initialize OpenAI Client
#     try:
#         client = OpenAI(api_key=api_key, timeout=90.0) # Set a default timeout
#     except Exception as e:
#         logging.error(f"Failed to initialize OpenAI client: {e}")
#         return

#     # 2. Extract Text
#     logging.info("--- Step 1: Extracting Text ---")
#     text_extract_start = time.time()
#     resume_text = get_resume_text(input_path)
#     text_extract_end = time.time()
#     if not resume_text:
#         logging.error(f"Failed to extract text from {input_path}. Aborting.")
#         return
#     logging.info(f"Text extraction completed in {text_extract_end - text_extract_start:.2f} seconds.")

#     # 3. Parse with LLM
#     logging.info("--- Step 2: Parsing with LLM ---")
#     parsed_data, llm_processing_time = parse_resume_with_llm(resume_text, client, model=args.model)
#     if not parsed_data:
#         logging.error(f"Failed to parse resume content using LLM for {input_path}. Aborting.")
#         return
#     logging.info(f"LLM parsing step completed in {llm_processing_time:.2f} seconds (includes API time and retries).")


#     # 4. Save Output
#     logging.info("--- Step 3: Saving Output ---")
#     save_start_time = time.time()
#     success = save_parsed_data(parsed_data, output_path)
#     save_end_time = time.time()

#     overall_end_time = time.time()
#     total_time = overall_end_time - overall_start_time

#     if success:
#         logging.info(f"Saving completed in {save_end_time - save_start_time:.2f} seconds.")
#         logging.info(f"--- Resume parsing process completed successfully in {total_time:.2f} seconds ---")
#     else:
#         logging.error(f"--- Resume parsing process failed during saving. Total time: {total_time:.2f} seconds ---")

# if __name__ == "__main__":
#     main()
